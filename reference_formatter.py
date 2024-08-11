import os
import re
import requests
import locale
from isbnlib import meta
from habanero import Crossref
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime

# Configura a localidade para português do Brasil, o que afeta a formatação de datas.
locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')

def set_document_styles(doc, lang):
    # Configura as margens do documento
    section = doc.sections[0]
    section.top_margin = Pt(85)  # 3 cm
    section.right_margin = Pt(85)  # 3 cm
    section.bottom_margin = Pt(57)  # 2 cm
    section.left_margin = Pt(57)  # 2 cm

    # Configura o estilo de fonte e parágrafo
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    style.paragraph_format.space_after = Pt(12)  # Espaço entre parágrafos
    style.paragraph_format.line_spacing = 1.0  # Espaço simples

    # Define o idioma dos parágrafos
    for paragraph in doc.paragraphs:
        paragraph_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        paragraph_run._element.rPr.lang = qn(lang)

def extract_reference_parts(reference):
    # Extrai autores, ano, título e outras informações de uma referência
    authors_match = re.match(r'^(.*?)(\.\s)', reference)
    authors = authors_match.group(1).strip() if authors_match else ''
    
    year_match = re.search(r'\b(\d{4})\b', reference)
    year = year_match.group(1) if year_match else ''
    
    title_match = re.search(r'\.\s(.*?)\.\s', reference)
    title = title_match.group(1).strip() if title_match else reference.split(".")[1].strip() if len(reference.split(".")) > 1 else ''
    
    additional_info = reference.split(title)[-1].strip() if title else reference.strip()
    additional_info = additional_info.strip('.').strip()
    
    city, publisher, extracted_year = extract_city_publisher_year(additional_info, year)
    additional_info = re.sub(r'\b(\d{4})\b', '', additional_info).strip(', ').strip()
    if city:
        additional_info = additional_info.replace(city, '').strip(', ').strip()
    if publisher:
        additional_info = additional_info.replace(publisher, '').strip(', ').strip()
    
    return authors, extracted_year, title, additional_info, city, publisher

def extract_city_publisher_year(additional_info, year):
    # Extrai cidade, editora e ano das informações adicionais da referência
    city = ''
    publisher = ''
    extracted_year = year

    # Verifica se o ano está presente nas informações adicionais
    if re.search(r'\b\d{4}\b', additional_info):
        extracted_year_match = re.search(r'\b(\d{4})\b', additional_info)
        extracted_year = extracted_year_match.group(1) if extracted_year_match else year
        additional_info = additional_info.replace(extracted_year, '').strip(', ').strip()

    # Extrai cidade e editora usando delimitadores ":" ou ","
    if ':' in additional_info:
        city_publisher = additional_info.split(':')
        city = city_publisher[0].strip()
        publisher = city_publisher[1].strip()
    elif ',' in additional_info:
        city_publisher = additional_info.split(',')
        city = city_publisher[0].strip()
        publisher = city_publisher[1].strip()
    else:
        publisher = additional_info.strip()

    return city, publisher, extracted_year

def bold_text(run, text):
    # Adiciona texto em negrito a uma run do parágrafo
    run.add_run(text).bold = True

def add_formatted_reference(paragraph, reference_parts):
    # Adiciona partes formatadas de uma referência a um parágrafo
    for part in reference_parts:
        if isinstance(part, dict) and part.get("bold"):
            bold_text(paragraph, part.get("text", ""))
        else:
            paragraph.add_run(part if isinstance(part, str) else part.get("text", ""))

def format_author_abnt(authors, ref_type=None):
    # Formata o nome dos autores no estilo ABNT
    names = authors.split(';')
    formatted_names = []

    for name in names:
        name = re.sub(r'[^\w\s]', '', name).strip()  # Remove caracteres especiais
        parts = name.split()

        # Remove todas as iniciais que são uma única letra
        parts = [part for part in parts if len(part) > 1]

        if ref_type == "isbn" and len(parts) > 1:
            last_name = parts[-1].upper().strip()
            first_names = ' '.join(parts[:-1]).capitalize().strip()
            if first_names:
                formatted_name = f"{last_name}, {first_names}"
            else:
                formatted_name = last_name
        else:
            if len(parts) > 1:
                last_name = parts[0].upper().strip()
                first_names = ' '.join(parts[1:]).capitalize().strip()
                if first_names:
                    formatted_name = f"{last_name}, {first_names}"
                else:
                    formatted_name = last_name
            else:
                formatted_name = name.upper()

        formatted_names.append(formatted_name)

    final_result = '; '.join(formatted_names).replace(",.", ".")
    return final_result + '.'

def format_author_apa7(authors, ref_type=None):
    # Formata o nome dos autores no estilo APA 7ª edição
    names = authors.split(';')
    formatted_names = []

    for name in names:
        name = re.sub(r'[^\w\s]', '', name).strip()  # Remove caracteres indesejados
        parts = name.split()

        if ref_type == "isbn" and len(parts) > 1:  # Condicional específica para ISBN
            last_name = parts[-1].capitalize().strip()
            first_names = ' '.join(parts[:-1]).capitalize().strip()
            if first_names:
                formatted_name = f"{last_name}, {first_names[0].upper()}."
            else:
                formatted_name = f"{last_name}."
        else:
            last_name = parts[0].capitalize().strip()
            first_names = ' '.join(parts[1:]).capitalize().strip()
            if first_names:
                formatted_name = f"{last_name}, {first_names[0].upper()}."
            else:
                formatted_name = f"{last_name}."

        formatted_names.append(formatted_name.strip())

    final_result = ', '.join(formatted_names)
    return final_result


def format_author_apa7(authors, ref_type=None):
    # Formata o nome dos autores no estilo APA 7ª edição, tratando nomes e iniciais
    names = authors.split(';')
    formatted_names = []

    for name in names:
        name = re.sub(r'[^\w\s]', '', name).strip()
        parts = name.split()

        if ref_type == "isbn" and len(parts) > 1:
            last_name = parts[-1].capitalize().strip()
            first_names = ' '.join(parts[:-1]).capitalize().strip()
            if first_names:
                formatted_name = f"{last_name}, {first_names[0].upper()}."
            else:
                formatted_name = f"{last_name}."
        else:
            last_name = parts[0].capitalize()
            first_names = ' '.join(parts[1:]).strip()
            formatted_name = f"{last_name}, {first_names[0].upper()}." if first_names else f"{last_name}."
        formatted_names.append(formatted_name.strip())

    final_result = ', '.join(formatted_names)
    final_result = re.sub(r',\s,', ',', final_result)  # Corrige possíveis vírgulas duplicadas
    return final_result

def process_reference_abnt(authors, title, city, publisher, year, additional_info=None, volume=None, number=None, pages=None, doi=None, url=None, publication_name=None):
    # Monta a referência formatada no estilo ABNT
    reference_parts = []
    
    reference_parts.append(f"{authors} ")  # Adiciona os autores

    if publication_name:  # Negritar o nome da publicação
        title = title.replace(" - ", ": ")  # Substituir hífen por dois pontos
        reference_parts.append(f"{title}. ")
        reference_parts.append({"text": publication_name, "bold": True})
        reference_parts.append(". ")
    else:  # Negritar o título principal se não houver publicação
        if ":" in title:
            main_title, subtitle = title.split(":", 1)
            reference_parts.append({"text": main_title.strip(), "bold": True})
            reference_parts.append(f": {subtitle.strip()}. ")
        else:
            reference_parts.append({"text": title, "bold": True})
            reference_parts.append(". ")
    
    if city and publisher:
        reference_parts.append(f"{city}: {publisher}, {year}. ")
    elif publisher:
        reference_parts.append(f"{publisher}, {year}. ")

    if additional_info:
        reference_parts.append(f"{additional_info} ")
    if volume:
        reference_parts.append(f"v. {volume} ")
    if number:
        reference_parts.append(f"n. {number} ")
    if pages:
        reference_parts.append(f"p. {pages}.")

    if doi:
        reference_parts.append(f" {year}. DOI: {doi}. ")
    if url:
        acesso = datetime.now().strftime('%d de %B de %Y')
        reference_parts.append(f"Disponível em: {url}. Acesso em: {acesso}. ")

    return reference_parts

def process_reference_apa7(authors, title, year, publisher=None, volume=None, number=None, pages=None, doi=None, url=None, publication_name=None):
    # Monta a referência formatada no estilo APA 7ª edição
    reference_parts = []
    
    reference_parts.append(f"{authors} ({year}). ")

    if publication_name:  # Negritar o nome da publicação
        title = title.replace(" - ", ": ")  # Substituir hífen por dois pontos
        reference_parts.append(f"{title}. ")
        reference_parts.append({"text": publication_name, "bold": True})
        reference_parts.append(". ")
    else:  # Negritar o título principal se não houver publicação
        if ":" in title:
            main_title, subtitle = title.split(":", 1)
            reference_parts.append({"text": main_title.strip(), "bold": True})
            reference_parts.append(f": {subtitle.strip()}. ")
        else:
            reference_parts.append({"text": title, "bold": True})
            reference_parts.append(". ")

    if publisher:
        reference_parts.append(f"{publisher}. ")

    if volume and number and pages:
        reference_parts.append(f"{volume}({number}), {pages}. ")

    if doi:
        reference_parts.append(f"https://doi.org/{doi} ")
    elif url:
        reference_parts.append(f"Disponível em: {url} ")

    return reference_parts

def process_reference(reference, ref_type="manual", identifier=None):
    # Processa uma referência de acordo com o tipo (manual, DOI ou ISBN)
    try:
        if ref_type == "doi":
            # Processa referência usando DOI
            data = get_data_by_doi(identifier)
            if data is None:
                raise ValueError(f"Erro ao buscar DOI {identifier}")
            authors = '; '.join([f"{auth['family']}, {auth['given'][0]}" for auth in data.get('author', [])])
            title = data.get('title', [''])[0]
            publication_year = str(data.get('issued', {}).get('date-parts', [[None]])[0][0])
            volume = data.get('volume', '')
            number = data.get('issue', '')
            pages = data.get('page', '')
            doi = identifier.split('https://doi.org/')[-1]
            publication_name = data.get('container-title', [''])[0]  # Nome da publicação

            authors_abnt = format_author_abnt(authors, ref_type="doi")
            authors_apa7 = format_author_apa7(authors, ref_type="doi")

            formatted_abnt = process_reference_abnt(authors_abnt, title, '', '', publication_year, volume=volume, number=number, pages=pages, doi=doi, publication_name=publication_name, url=data.get('URL'))
            formatted_apa7 = process_reference_apa7(authors_apa7, title, publication_year, volume=volume, number=number, pages=pages, doi=doi, publication_name=publication_name, url=data.get('URL'))
        
        elif ref_type == "isbn":
            # Processa referência usando ISBN
            data = get_data_by_isbn(identifier)
            if data is None:
                raise ValueError(f"Erro ao buscar ISBN {identifier}")
            authors = data.get('Authors', [''])[0]
            title = data.get('Title', '')
            publication_year = data.get('Year', '')
            additional_info = data.get('Publisher', 'Informação não encontrada')
            city, publisher, extracted_year = extract_city_publisher_year(additional_info, publication_year)

            authors_abnt = format_author_abnt(authors, ref_type="isbn")
            authors_apa7 = format_author_apa7(authors, ref_type="isbn")

            formatted_abnt = process_reference_abnt(authors_abnt, title, city, publisher, extracted_year)
            formatted_apa7 = process_reference_apa7(authors_apa7, title, extracted_year, publisher=publisher)
        
        else:
            # Processa referência manualmente
            authors, extracted_year, title, additional_info, city, publisher = extract_reference_parts(reference)
            authors_abnt = format_author_abnt(authors)
            authors_apa7 = format_author_apa7(authors)

            formatted_abnt = process_reference_abnt(authors_abnt, title, city, publisher, extracted_year, additional_info=additional_info)
            formatted_apa7 = process_reference_apa7(authors_apa7, title, extracted_year, publisher=publisher)
        
        print(f"Referência processada (ABNT): {''.join([str(part) for part in formatted_abnt])}")  # Depuração
        print(f"Referência processada (APA7): {''.join([str(part) for part in formatted_apa7])}")  # Depuração
        
        return formatted_abnt, formatted_apa7
    
    except Exception as e:
        log_error(f"Erro ao processar referência: {reference}. Erro: {str(e)}")
        return None, None

def get_data_by_doi(doi):
    # Busca dados de referência usando DOI via Crossref
    cr = Crossref()
    try:
        if doi.startswith('https://doi.org/'):
            doi = doi.split('https://doi.org/')[-1]
        data = cr.works(ids=doi)
        return data['message']
    except requests.exceptions.HTTPError as e:
        log_error(f"Erro ao buscar DOI {doi}: {e}")
        return None

def get_data_by_isbn(isbn):
    # Busca dados de referência usando ISBN via isbnlib
    try:
        data = meta(isbn)
        if data is not None and data.get('ISBN') and isbn not in data.get('ISBN'):
            raise ValueError(f"isbn request != isbn response ({isbn} not in {data.get('ISBN')})")
        return data
    except Exception as e:
        log_error(f"Erro ao buscar ISBN {isbn}: {e}")
        return None

def log_error(error_message):
    # Registra mensagens de erro em um arquivo de log
    script_dir = os.path.dirname(os.path.abspath(__file__))
    log_file = os.path.join(script_dir, 'error_log.txt')
    with open(log_file, 'a', encoding='utf-8') as log:
        log.write(f"{error_message}\n")

def save_references_abnt(references, file_name):
    # Salva as referências formatadas em ABNT em um arquivo .docx
    doc = Document()
    set_document_styles(doc, lang='pt-BR')
    
    doc.add_heading('REFERÊNCIAS', level=1).bold = True
    
    for reference_parts in references:
        p = doc.add_paragraph()
        p.alignment = 3  # Justificado
        add_formatted_reference(p, reference_parts)
    
    doc.save(file_name)

def save_references_apa7(references, file_name):
    # Salva as referências formatadas em APA 7ª edição em um arquivo .docx
    doc = Document()
    set_document_styles(doc, lang='en-US')
    
    doc.add_heading('REFERENCES', level=1).bold = True
    
    for reference_parts in references:
        p = doc.add_paragraph()
        p.alignment = 3  # Justificado
        add_formatted_reference(p, reference_parts)
    
    doc.save(file_name)

def process_references_from_file(input_file):
    # Processa referências de um arquivo de texto
    if not os.path.exists(input_file):
        print(f"Erro: o arquivo {input_file} não foi encontrado.")
        return
    
    abnt_references = []
    apa7_references = []

    # Lê as referências do arquivo
    with open(input_file, 'r', encoding='utf-8') as file:
        references = file.readlines()

    for reference in references:
        reference = reference.strip()
        if reference.startswith("10.") or reference.startswith("https://doi.org/"):
            abnt, apa7 = process_reference(reference, ref_type="doi", identifier=reference)
        elif reference.startswith("978") or reference.startswith("isbn"):
            abnt, apa7 = process_reference(reference, ref_type="isbn", identifier=reference)
        else:
            abnt, apa7 = process_reference(reference)
        
        if abnt and apa7:
            abnt_references.append(abnt)
            apa7_references.append(apa7)

    # Salva as referências processadas nos arquivos .docx
    script_dir = os.path.dirname(os.path.abspath(__file__))
    save_references_abnt(abnt_references, os.path.join(script_dir, 'referencias_abnt.docx'))
    save_references_apa7(apa7_references, os.path.join(script_dir, 'referencias_apa7.docx'))

# Diretório do script e arquivo de entrada
script_dir = os.path.dirname(os.path.abspath(__file__))
input_file = os.path.join(script_dir, 'referencias.txt')

# Processa as referências do arquivo
process_references_from_file(input_file)
